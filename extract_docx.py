#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
提取Word文档内容的脚本
"""
import sys
import io

# 设置标准输出为UTF-8编码
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

try:
    from docx import Document
except ImportError:
    print("错误: 需要安装python-docx库")
    print("请运行: pip install python-docx")
    sys.exit(1)

def extract_docx_content(file_path):
    """提取Word文档的所有文本内容"""
    try:
        doc = Document(file_path)
        
        print("=" * 80)
        print(f"文档: {file_path}")
        print("=" * 80)
        print()
        
        # 提取所有段落
        for i, paragraph in enumerate(doc.paragraphs, 1):
            # 清理文本，移除零宽字符等特殊字符
            text = paragraph.text.strip().replace('\u200b', '').replace('\u200c', '').replace('\u200d', '')
            if text:  # 只打印非空段落
                print(f"[段落 {i}]")
                print(text)
                print()
        
        # 提取表格内容
        if doc.tables:
            print("\n" + "=" * 80)
            print("表格内容:")
            print("=" * 80)
            for table_idx, table in enumerate(doc.tables, 1):
                print(f"\n[表格 {table_idx}]")
                for row_idx, row in enumerate(table.rows, 1):
                    row_data = [cell.text.strip() for cell in row.cells]
                    print(f"  行 {row_idx}: {' | '.join(row_data)}")
        
        # 将内容保存到文本文件
        output_file = file_path.replace('.docx', '_extracted.txt')
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"文档: {file_path}\n")
            f.write("=" * 80 + "\n\n")
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip().replace('\u200b', '').replace('\u200c', '').replace('\u200d', '')
                if text:
                    f.write(text + "\n\n")
            
            if doc.tables:
                f.write("\n" + "=" * 80 + "\n")
                f.write("表格内容:\n")
                f.write("=" * 80 + "\n")
                for table_idx, table in enumerate(doc.tables, 1):
                    f.write(f"\n[表格 {table_idx}]\n")
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        f.write(' | '.join(row_data) + "\n")
        
        print(f"\n内容已保存到: {output_file}")
        
    except Exception as e:
        print(f"错误: 无法读取文档 - {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    file_path = "prd.docx"
    extract_docx_content(file_path)

